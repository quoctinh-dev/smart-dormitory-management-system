import argparse
import os
try:
    from docx import Document
except ImportError:
    print("Vui lòng cài đặt thư viện bằng lệnh: pip install python-docx")
    exit(1)

def insert_text_after_heading(doc_path, heading_text, content_file, output_path):
    print(f"[*] Đang mở file Word: {doc_path}")
    if not os.path.exists(doc_path):
        print(f"[!] Lỗi: Không tìm thấy file {doc_path}. Vui lòng tạo sẵn một file Word.")
        return
        
    doc = Document(doc_path)
    
    if not os.path.exists(content_file):
        print(f"[!] Lỗi: Không tìm thấy file nội dung {content_file}.")
        return

    with open(content_file, 'r', encoding='utf-8') as f:
        content_to_insert = f.read()

    # Tìm đoạn văn (Paragraph) có chứa tiêu đề mong muốn
    heading_found = False
    for i, p in enumerate(doc.paragraphs):
        # Kiểm tra xem đây có phải là một Heading không và có chứa từ khóa không
        if p.style.name.startswith('Heading') and heading_text.lower() in p.text.lower():
            heading_found = True
            print(f"[*] Đã tìm thấy Heading: '{p.text}'. Đang chèn nội dung vào bên dưới...")
            
            # python-docx không hỗ trợ insert_after trực tiếp, nên ta phải insert_before ở đoạn tiếp theo
            if i + 1 < len(doc.paragraphs):
                new_p = doc.paragraphs[i+1].insert_paragraph_before(content_to_insert)
            else:
                doc.add_paragraph(content_to_insert)
            break
            
    if not heading_found:
        print(f"[!] Cảnh báo: Không tìm thấy Heading nào chứa chữ '{heading_text}' trong file Word.")
        print(f"[*] Sẽ chèn nội dung vào cuối file thay thế.")
        doc.add_paragraph(content_to_insert)
        
    doc.save(output_path)
    print(f"[OK] Cập nhật thành công! File mới được lưu tại: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Script tự động chèn Text/Markdown vào file Word (Dành cho Vỹ)")
    parser.add_argument("--input", required=True, help="Đường dẫn đến file text/markdown chứa nội dung")
    parser.add_argument("--target", required=True, help="Đường dẫn file Word gốc (thesis.docx)")
    parser.add_argument("--heading", required=True, help="Từ khóa của Heading muốn chèn chữ xuống dưới")
    parser.add_argument("--output", default="thesis_updated.docx", help="Đường dẫn file Word đầu ra")
    
    args = parser.parse_args()
    
    insert_text_after_heading(args.target, args.heading, args.input, args.output)
